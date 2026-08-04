import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os
import copy
from docx.table import _Row
from urllib.parse import quote
import base64

# --- CONFIGURAÇÕES ---
st.set_page_config(page_title="Gerador ATS - SSMA", layout="wide")

BASE_PATH = os.getcwd()

# Caminhos dos assets visuais
FUNDO_PATH = os.path.join(BASE_PATH, "fundo.jpg")
LOGO_PATH = os.path.join(BASE_PATH, "logo.png")

SHEET_ID = "1y98U3eK7JXJqQaMC0i7eFbwpvp97Nuyeml5Dis0UCUg"

# -----------------------------
# MAPEAMENTO DOS TEMPLATES POR EMPRESA
# -----------------------------
TEMPLATES_EMPRESAS = {
    "Benteler": "T.SHE.046 Safe Job Analyses Bentler.docx",
    "Macromaq": "T.SHE.046 Safe Job Analyses Macromaq.docx"
}

# --- LAYOUT E CSS ---
def get_base64(bin_file):
    if not os.path.exists(bin_file):
        return ""
    with open(bin_file, 'rb') as f:
        data = f.read()
    return base64.b64encode(data).decode()

def aplicar_layout():
    fundo = get_base64(FUNDO_PATH)
    logo = get_base64(LOGO_PATH)
    
    if fundo:
        estilo_fundo = f"""
        .stApp {{
            background-image: linear-gradient(rgba(0,0,0,0.5), rgba(0,0,0,0.5)), url("data:image/jpeg;base64,{fundo}") !important;
            background-size: cover !important;
            background-position: center center !important;
            background-attachment: fixed !important;
        }}
        """
    else:
        estilo_fundo = ".stApp {{ background-color: #1e293b; }}"

    st.markdown(f"""
    <style>
    /* REMOVER BARRA LATERAL E NAVEGAÇÃO NATIVA */
    [data-testid="stSidebar"], [data-testid="stSidebarNav"] {{
        display: none;
    }}

    {estilo_fundo}
    
    .stApp > header {{
        background-color: transparent !important;
    }}

    .stSelectbox label, .stTextInput label, div[data-testid="stCheckbox"] label p {{
        color: white !important;
        background: rgba(0,0,0,0.7);
        padding: 5px 12px;
        border-radius: 8px;
        font-weight: bold;
    }}
    .stButton > button {{
        background: #2c3e50;
        color: #f9cc0b;
        border: 2px solid #f9cc0b;
        border-radius: 10px;
        height: 55px;
        font-weight: bold;
        width: 100%;
        font-size: 18px;
    }}
    .header-container {{
        display: flex;
        align-items: center;
        background: rgba(255,255,255,0.95);
        padding: 20px;
        border-radius: 15px;
        margin-bottom: 30px;
    }}
    .footer {{
        position: fixed;
        left: 0;
        bottom: 0;
        width: 100%;
        background: rgba(0,0,0,0.8);
        color: white;
        text-align: center;
        padding: 10px;
        font-size: 13px;
        z-index: 999;
    }}
    </style>
    <div class="header-container">
        <img src="data:image/png;base64,{logo}" width="320">
        <h1 style="margin-left:25px;color:#2c3e50;">Gerador ATS - SSMA</h1>
    </div>
    """, unsafe_allow_html=True)

aplicar_layout()

# Botão para limpar cache e atualizar dados da planilha instantaneamente
if st.button("🔄 Atualizar Dados da Planilha Agora"):
    st.cache_data.clear()
    st.rerun()

# -----------------------------
# LEITURA DA ABA "Banco_APRs" DO GOOGLE SHEETS
# -----------------------------
@st.cache_data(ttl=300)
def carregar_dados_sheets():
    try:
        url = f"https://docs.google.com/spreadsheets/d/{SHEET_ID}/gviz/tq?tqx=out:csv&sheet={quote('Banco_APRs'.strip())}"
        df = pd.read_csv(url, dtype=str)
        df.columns = df.columns.str.lower().str.strip()
        df.rename(columns={
            'ações': 'acoes', 'acao': 'acoes',
            'atividade': 'atividade', 'tarefa': 'tarefa',
            'risco': 'risco', 'área': 'area'
        }, inplace=True)
        return df
    except Exception as e:
        return pd.DataFrame()

df_excel = carregar_dados_sheets()
LISTA_ATIVIDADES = df_excel["atividade"].dropna().unique().tolist() if not df_excel.empty else []


# -----------------------------
# FUNÇÃO AUXILIAR DE SUBSTITUIÇÃO
# -----------------------------
def substituir_texto_paragrafo(p, mapeamento):
    texto_completo = "".join([run.text for run in p.runs])
    houve_mudanca = False
    
    for tag, valor in mapeamento.items():
        if tag in texto_completo:
            texto_completo = texto_completo.replace(tag, str(valor))
            houve_mudanca = True
            
    if houve_mudanca:
        if p.runs:
            p.runs[0].text = texto_completo
            for run in p.runs[1:]:
                p._p.remove(run._r)
        else:
            p.add_run(texto_completo)


# -----------------------------
# PROCESSAMENTO DO DOCX
# -----------------------------
def substituir_docx(doc, mapeamento, passos_atividade):
    for p in doc.paragraphs:
        substituir_texto_paragrafo(p, mapeamento)

    for table in doc.tables:
        linha_modelo = None
        idx_tarefa, idx_risco, idx_acao = None, None, None
        
        for row in table.rows:
            encontrou_na_linha = False
            for idx, cell in enumerate(row.cells):
                texto_celula = cell.text.upper()
                if "{{TAREFAS}}" in texto_celula:
                    linha_modelo, idx_tarefa = row, idx
                    encontrou_na_linha = True
                if "{{RISCOS}}" in texto_celula:
                    linha_modelo, idx_risco = row, idx
                    encontrou_na_linha = True
                if "{{ACOES}}" in texto_celula:
                    linha_modelo, idx_acao = row, idx
                    encontrou_na_linha = True
            if encontrou_na_linha:
                break
        
        if linha_modelo is not None:
            passos_lista = passos_atividade.to_dict(orient="records")
            linha_referencia = linha_modelo
            
            for i, item in enumerate(passos_lista):
                tarefa = "" if pd.isna(item.get("tarefa")) else str(item["tarefa"]).strip()
                risco = "" if pd.isna(item.get("risco")) else str(item["risco"]).strip()
                acao = "" if pd.isna(item.get("acoes")) else str(item["acoes"]).strip()

                if i == 0:
                    nova_linha = linha_modelo
                else:
                    tr_copiado = copy.deepcopy(linha_modelo._tr)
                    linha_referencia._tr.addnext(tr_copiado)
                    nova_linha = _Row(tr_copiado, table)
                    linha_referencia = nova_linha
                
                def preencher_celula_limpa(celula, texto):
                    if not celula.paragraphs:
                        celula.add_paragraph()
                    p = celula.paragraphs[0]
                    p.text = ""
                    for extra_p in celula.paragraphs[1:]:
                        p_xml = extra_p._p
                        p_xml.getparent().remove(p_xml)
                    
                    run = p.add_run(texto)
                    run.font.size = Pt(7)
                
                if idx_tarefa is not None:
                    preencher_celula_limpa(nova_linha.cells[idx_tarefa], f"{i+1}. {tarefa}" if tarefa else "")
                if idx_risco is not None:
                    preencher_celula_limpa(nova_linha.cells[idx_risco], f"• {risco}" if risco else "")
                if idx_acao is not None:
                    preencher_celula_limpa(nova_linha.cells[idx_acao], f"• {acao}" if acao else "")
                    
        else:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        substituir_texto_paragrafo(p, mapeamento)


# -----------------------------
# GERADOR PRINCIPAL
# -----------------------------
def gerar_documento(dados, atividade_selecionada, passos, caminho_template):
    if not os.path.exists(caminho_template):
        st.error(f"Template não encontrado no caminho: {caminho_template}")
        return None

    doc = Document(caminho_template)

    mapeamento = {
        "{{CONTRATADA}}": dados["contratada"].upper(),
        "{{LOCAL}}": dados["local"].upper(),
        "{{AREA}}": dados["area"].upper(),
        "{{ÁREA}}": dados["area"].upper(),
        "{{ATIVIDADE}}": atividade_selecionada.upper(),
        "{{PROCESSO}}": dados["processo"].upper()
    }

    substituir_docx(doc, mapeamento, passos)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# -----------------------------
# INTERFACE STREAMLIT
# -----------------------------
st.subheader("📋 Configuração do Template")
empresa_selecionada = st.selectbox("Selecione a Empresa cliente (Template):", list(TEMPLATES_EMPRESAS.keys()))
caminho_template_escolhido = TEMPLATES_EMPRESAS[empresa_selecionada]

st.divider()

col1, col2 = st.columns(2)

with col1:
    atividade = st.selectbox("Atividade", LISTA_ATIVIDADES) if not df_excel.empty else st.selectbox("Atividade", ["Nenhuma atividade carregada"])
    processo = st.text_input("Processo", "MANUTENÇÃO MECÂNICA")

with col2:
    contratada = st.text_input("Contratada", empresa_selecionada.upper())
    local = st.text_input("Local", "OFICINA")
    area = st.text_input("Área", "EMPILHADEIRA")

passos = df_excel[df_excel["atividade"] == atividade] if not df_excel.empty else pd.DataFrame()

st.subheader("Prévia dos dados da aba 'Banco_APRs'")
st.dataframe(passos, use_container_width=True)

if st.button("🚀 Gerar Documento", type="primary"):
    if df_excel.empty:
        st.error("A aba 'Banco_APRs' não foi carregada corretamente do Google Sheets. Verifique o compartilhamento da planilha.")
    elif passos.empty:
        st.warning("Nenhum passo encontrado para a atividade selecionada.")
    else:
        dados = {"contratada": contratada, "local": local, "area": area, "processo": processo}
        
        resultado = gerar_documento(dados, atividade, passos, caminho_template_escolhido)

        if resultado:
            st.success(f"Documento gerado com sucesso utilizando o template da {empresa_selecionada}!")
            st.download_button(
                "📥 Baixar DOCX",
                resultado,
                file_name=f"ATS_{empresa_selecionada}_{atividade.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# --- FOOTER ---
st.markdown("""<div class="footer">© 2026 Gestão Documentos | Desenvolvido por: Dilceu Junior</div>""", unsafe_allow_html=True)
