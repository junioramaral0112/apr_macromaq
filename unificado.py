import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.table import _Row
import io
import os
import copy
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.drawing.image import Image as OpenpyxlImage
from datetime import datetime
import re

# ==============================================================================
# CONFIGURAÇÃO DA PÁGINA
# ==============================================================================
st.set_page_config(
    page_title="Sistema Integrado de Segurança - SSMA",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==============================================================================
# BANCO DE DADOS DE EQUIPAMENTOS (Checklist)
# ==============================================================================
BANCO_EQUIPAMENTOS = {
    "Compressor de Ar": [
        "O purgador na parte inferior do reservatório foi aberto para drenar totalmente a água acumulada?",
        "O manômetro indica a pressão interna de forma limpa, clara e possui marcação de faixa crítica?",
        "A grade metálica de proteção da correia e das polias está instalada, firme e impede o acesso de membros?",
        "O nível de óleo do bloco compressor está na marca correta do visor (risco de travamento e explosão)?",
        "A válvula de segurança foi testada manualmente e está desimpedida e operacional (NR-13)?",
        "O pressostato está desarmando o motor automaticamente ao atingir a pressão máxima de trabalho?",
        "A fiação elétrica de alimentação está protegida por conduítes e longe de fontes de calor ou umidade?",
        "A carcaça do motor elétrico e o corpo do compressor possuem aterramento elétrico visível e íntegro?",
        "O equipamento está instalado em local ventilado e com espaço seguro para manutenção?",
        "O registro geral de saída de ar está operando perfeitamente e sem vazamentos nas conexões?"
    ],
    "Extensão Elétrica": [
        "O cabo elétrico está totalmente desenrolado durante o uso para evitar o efeito bobina e superaquecimento?",
        "A capa isolante externa do cabo está sem cortes, esmagamentos, emendas com fita ou fios aparentes?",
        "O plugue macho possui todos os pinos (incluindo o terra), sem trincas na carcaça e sem pinos tortos?",
        "O bloco de tomadas fêmea está limpo, sem sinais de derretimento, queimados ou folga nos contatos?",
        "O cabo possui o pino de aterramento functional conectado nas duas extremidades (NR-10)?",
        "A extensão é adequada para a potência do equipamento que será ligado (corrente/amperagem compatível)?",
        "O cabo está posicionado fora de rotas de passagem de pessoas e veículos para evitar tropeços e prensagem?",
        "A extensão está completamente livre de contato com poças de água, umidade ou superfícies metálicas vivas?"
    ],
    "Furadeira a Bateria": [
        "O mandril está apertando a broca firmemente com a chave/sistema manual e sem apresentar folgas no eixo?",
        "O gatilho de acionamento possui sistema de controle de velocidade e desliga o motor imediatamente ao ser solto?",
        "A bateria se encaixa e trava perfeitamente no corpo da ferramenta, sem folgas ou risco de queda?",
        "A carcaça plástica externa está totalmente inteira, livre de trincas, quebras ou partes soltas?",
        "O botão seletor de reversão de rotação funciona corretamente e não trava no meio do curso?",
        "A carcaça da bateria está sem estufamentos, vazamentos de fluido químico ou oxidação nos contatos?",
        "A ferramenta possui ponto para fixação de fita de segurança antiqueda (se utilizada em trabalhos em altura)?",
        "O punho emborrachado está firme, limpo, sem resíduos de óleo ou graxa que possam causar deslizamento?"
    ],
    "Furadeira Elétrica": [
        "O cabo elétrico de alimentação está totalmente íntegro, sem emendas, cortes na isolação ou fios expostos?",
        "O plugue de tomada está em perfeitas condições, sem pinos tortos, rachaduras ou sinais de sobreaquecimento?",
        "O mandril está perfeitamente alinhado (sem jogo) e a broca fica travada com aperto correto?",
        "O gatilho de acionamento liga e desliga instantaneamente, e o botão de trava de gatilho desarma ao menor toque?",
        "O punho auxiliar regulável está instalado, firme e posicionado para garantir a empunhadura segura?",
        "Ausência de faíscas excessivas no coletor/escovas ou cheiro de queimado no motor durante a operação?",
        "As aberturas de ventilação do motor estão totalmente desimpedidas e limpas de poeiras ou cavacos?",
        "A chave de aperto do mandril foi retirada do bocal antes de ligar a ferramenta (risco de projeção)?"
    ],
    "Lavadora de Alta Pressão": [
        "A mangueira de alta pressão está sem dobras estruturais, rachaduras, bolhas ou malha de aço exposta?",
        "O cabo elétrico de alimentação and o plugue estão secos, sem emendas e com isolamento duplo íntegro?",
        "Os engates rápidos e conexões da mangueira de alta pressão e entrada de água estão firmes e sem vazamentos?",
        "O gatilho da pistola possui trava física de segurança contra acionamento acidental funcionando?",
        "O fluxo de água corta imediatamente quando o gatilho da pistola é liberado pelo operador?",
        "A lança, ponteira e o bico regulador estão desobstruídos e rosqueados firmemente até o fim?",
        "O equipamento possui disjuntor interno ou dispositivo DR integrado no cabo para proteção contra choques?",
        "O motor elétrico desliga automaticamente ao soltar o gatilho (sistema Total Stop operacional)?"
    ],
    "Macaco Hidráulico Garrafa": [
        "O corpo de aço e a base de sustentação do macaco estão sem deformações, trincas na solda ou fissuras?",
        "O nível de fluido hidráulico está correto e a haste eleva suavemente até o curso limite sem travar?",
        "Ausência total de gotejamentos ou vazamentos de óleo na base, gaxetas, pistão e na válvula de alívio?",
        "O fuso extensor roscado interno sobe e desce livremente com as mãos e não apresenta fios de rosca gastos?",
        "A válvula de alívio retém a pressão estática da carga sem permitir microdescidas ou perda de altura?",
        "O prato superior de apoio possui ranhuras antiderrapantes limpas, profundas e sem desgastes metálicos?",
        "A alavanca metálica de acionamento está perfeitamente reta, sem trincas ou improvisações na estrutura?",
        "A capacidade máxima de carga nominal está gravada de forma perfeitamente visível e legível no corpo do equipamento?"
    ],
    "Macaco Hidráulico Jacaré": [
        "O nível de óleo do reservatório está correto e garante a elevação máxima sem perda de curso ou cabeçadas?",
        "Ausência de vazamentos de óleo sob o chassi, nos cilindros hidráulicos, conexões e gaxetas?",
        "As rodas dianteiras e os rodízios traseiros giram e direcionam perfeitamente, sem travamentos?",
        "O prato de apoio de carga está centralizado no eixo, com a borracha/ranhura antiderrapante em bom estado?",
        "A alavanca de acionamento encaixa de forma firme no suporte mecânico e o pino de trava está seguro?",
        "A válvula de alívio/retorno abre de forma gradual permitindo o controle preciso e suave da descida da carga?",
        "As articulações e braços elevadores mecânicos estão sem folgas excessivas, empenamentos ou trincas?",
        "O chassi estrutural de aço está alinhado, plano e sem torções que comprometam a estabilidade do conjunto?"
    ],
    "Parafusadeira": [
        "O mandril de engate rápido ou porta-bits está fixando os acessórios firmemente e sem risco de soltura?",
        "O anel seletor de regulagem de torque gira suavemente e trava com precisão nas posições escolheras?",
        "O gatilho liga/desliga e o seletor mecânico de sentido de rotação funcionam sem apresentar travamentos?",
        "A carcaça plástica externa e os revestimentos emborrachados estão secos e livres de óleo ou graxas?",
        "Os cabos de alimentação elétrica ou as travas da bateria estão em perfeitas condições de conservação?",
        "O sistema de iluminação a LED auxiliar (se houver) está acendendo para garantir visibilidade da tarefa?",
        "Ausência de ruídos internos anormais na caixa de engrenagens ou vibração excessiva nas mãos do operador?"
    ],
    "Prensa Hidráulica": [
        "O manômetro indicador de pressão operacional está limpo, visível, calibrado e com lacre inviolado?",
        "Ausência completa de vazamentos de óleo em todas as mangueiras flexíveis, conexões rígidas e cilindro principal?",
        "A estrutura metálica do pórtico (chassi), a mesa móvel e os pinos de sustentação estão sem deformações ou trincas?",
        "A tela/grade de proteção frontal/lateral regulável contra estilhaços está instalada, firme e operacional (NR-12)?",
        "O comando de acionamento (alavanca, pedal ou bimano) responde prontamente e possui proteção contra toque acidental?",
        "O retorno do pistão hidráulico ocorre de forma suave, contínua e completa ao aliviar a pressão do sistema?",
        "Os blocos de apoio em 'V' ou calços da mesa estão sem trincas, rebarbas e assentam-se de forma plana?",
        "A área ao redor da prensa possui demarcação de piso e está totalmente desimpedida de materiais e óleo?"
    ],
    "Talha": [
        "O cabo de aço ou corrente de carga está livre de deformações, nós, amassados, corrosão ou elos esticados?",
        "O gancho de carga possui trava de segurança mecânica com mola funcionando e fechando 100% (NR-11)?",
        "O moitão, a polia guia e a guia de corrente giram livremente, sem apresentar desgastes ou ruídos anormais?",
        "A estrutura de fixação superior (olhal, viga ou troley manual/elétrico) está firme, sem trincas ou folgas?",
        "O sistema de freio mecânico/automático retém a carga suspensa firmemente (teste de retenção diário)?",
        "A botoeira de comando (se elétrica) possui cabo de sustentação de aço externo para não forçar os fios elétricos?",
        "O botão de parada de emergência (botão soco) na botoeira desliga o sistema imediatamente quando acionado?",
        "Os limitadores elétricos ou mecânicos de fim de curso (superior e inferior) atuam perfeitamente?",
        "A etiqueta indicativa da Capacidade Máxima de Carga (CMU) está fixada e perfeitamente legível à distância?",
        "A corrente de acionamento manual (se houver) não apresenta elos abertos ou risco de desengate da polia?"
    ]
}

# ==============================================================================
# FUNÇÕES DO CHECKLIST (Excel) - COM IMAGENS
# ==============================================================================
def gerar_excel_checklist(nome_equipamento, itens_checagem, caminho_salvamento=None):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"Checklist - {nome_equipamento[:15]}"
    
    # Configuração de Página
    ws.page_setup.orientation = ws.ORIENTATION_PORTRAIT
    ws.page_setup.paperSize = ws.PAPERSIZE_A4
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 1
    ws.views.sheetView[0].showGridLines = True

    # Estilos
    fina = Side(border_style="thin", color="000000")
    borda_grade = Border(left=fina, right=fina, top=fina, bottom=fina)
    fill_tabela_header = PatternFill(start_color="D9D9D9", fill_type="solid")
    
    font_titulo = Font(name="Arial", size=14, bold=True)
    font_sub = Font(name="Arial", size=10, bold=True)
    font_corpo = Font(name="Arial", size=9)              
    font_bold_corpo = Font(name="Arial", size=9, bold=True)
    font_legenda = Font(name="Arial", size=11, bold=True)

    # 1. CABEÇALHO SUPERIOR
    ws.merge_cells("A1:AB2")
    
    # Adicionar Logo Heineken
    caminho_logo = r"C:\Users\dilceu.gomes\Desktop\check list\heineken.png"
    if os.path.exists(caminho_logo):
        try:
            img_logo = OpenpyxlImage(caminho_logo)
            img_logo.width = 120
            img_logo.height = 60
            ws.add_image(img_logo, "A1")
        except Exception as e:
            st.warning(f"Não foi possível carregar a logo Heineken: {e}")
    
    # Texto do Título
    ws["A1"] = f"Checklist de Segurança: {nome_equipamento}"
    ws["A1"].font = font_titulo
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    
    # Adicionar Seta Verde
    caminho_seta = r"C:\Users\dilceu.gomes\Desktop\check list\seta.png"
    if os.path.exists(caminho_seta):
        try:
            img_seta = OpenpyxlImage(caminho_seta)
            img_seta.width = 50   
            img_seta.height = 50 
            ws.add_image(img_seta, "Z1")
        except Exception as e:
            st.warning(f"Não foi possível carregar a seta: {e}")

    # Blocos de Assinaturas
    ws["AC1"] = "SEGURANÇA"; ws["AE1"] = "LÍDER"; ws["AG1"] = "OPERADOR"
    ws.merge_cells("AC1:AD1"); ws.merge_cells("AE1:AF1"); ws.merge_cells("AG1:AH1")
    ws.merge_cells("AC2:AD2"); ws.merge_cells("AE2:AF2"); ws.merge_cells("AG2:AH2")

    # 2. DADOS DE IDENTIFICAÇÃO E EPIs
    ws.merge_cells("A3:M3"); ws["A3"] = "Empresa:"
    ws.merge_cells("A4:J4"); ws["A4"] = "Equip. Nº:"
    ws.merge_cells("K4:M4")
    ws.merge_cells("A5:M8") 
    ws.merge_cells("N3:T8")
    
    # EPIs
    ws.merge_cells("U3:AB3"); ws["U3"] = "Uso obrigatório dos EPIs:"
    
    epis = ["Capacete de segurança", "Protetor auricular", "Óculos de segurança", "Calçado de segurança", "Luva de proteção adequada"]
    for idx, epi in enumerate(epis, 4):
        ws.cell(row=idx, column=21, value="[  ]").alignment = Alignment(horizontal="center")
        ws.merge_cells(start_row=idx, start_column=22, end_row=idx, end_column=28) 
        ws.cell(row=idx, column=22, value=epi).font = font_corpo

    # Mês e SG
    ws.merge_cells("AC3:AH3"); ws["AC3"] = "Mês:        /"
    ws.merge_cells("AC4:AH8")
    
    # Adicionar Imagem SG
    caminho_sg = r"C:\Users\dilceu.gomes\Desktop\check list\SG.png"
    if os.path.exists(caminho_sg):
        try:
            img_sg = OpenpyxlImage(caminho_sg)
            img_sg.width = 120
            img_sg.height = 95
            ws.add_image(img_sg, "AC4")
        except Exception as e:
            st.warning(f"Não foi possível carregar a imagem SG: {e}")

    # Bordas no cabeçalho
    for r in range(1, 9):
        for c in range(1, 35):
            cell = ws.cell(row=r, column=c)
            cell.border = borda_grade
            if r in [1, 2] and c in [29, 31, 33]: 
                cell.font = Font(name="Arial", size=8, bold=True)
                cell.alignment = Alignment(horizontal="center", vertical="center")
            elif r < 5 and cell.value and c < 21:
                cell.font = font_sub

    # 3. TABELA PRINCIPAL
    headers = ["Item", "Itens de Segurança para Checar:", "Frequência"] + [str(i) for i in range(1, 32)]
    for c_idx, text in enumerate(headers, 1):
        cell = ws.cell(row=9, column=c_idx, value=text)
        cell.font = font_sub
        cell.fill = fill_tabela_header
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = borda_grade

    # 4. ITENS DO CHECKLIST
    row_atual = 10
    for item_num, texto in enumerate(itens_checagem, 1):
        ws.cell(row=row_atual, column=1, value=item_num).alignment = Alignment(horizontal="center", vertical="center")
        ws.cell(row=row_atual, column=3, value="Diário").alignment = Alignment(horizontal="center", vertical="center")
        cell_texto = ws.cell(row=row_atual, column=2, value=texto)
        cell_texto.font = font_corpo
        cell_texto.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)

        for c_idx in range(1, 35):
            c_cell = ws.cell(row=row_atual, column=c_idx)
            if c_idx != 2:
                c_cell.font = font_corpo
            c_cell.border = borda_grade
            
        ws.row_dimensions[row_atual].height = 32 if len(texto) > 60 else 22
        row_atual += 1

    # 5. VISTOS FINAIS
    vistos_finais = ["Visto de inspeção do operador do equipamento", "Visto do líder / encarregado / téc. segurança contratada"]
    for vf in vistos_finais:
        ws.merge_cells(start_row=row_atual, start_column=1, end_row=row_atual, end_column=2)
        ws.cell(row=row_atual, column=1, value=vf).font = font_sub
        ws.cell(row=row_atual, column=1).alignment = Alignment(horizontal="left", vertical="center")
        ws.cell(row=row_atual, column=3, value="Diário").font = font_sub
        ws.cell(row=row_atual, column=3).alignment = Alignment(horizontal="center", vertical="center")
        for c_idx in range(1, 35):
            ws.cell(row=row_atual, column=c_idx).border = borda_grade
        row_atual += 1

    # 6. OBSERVAÇÃO
    ws.merge_cells(start_row=row_atual, start_column=1, end_row=row_atual, end_column=34)
    cell_obs = ws.cell(row=row_atual, column=1, value="Obs: Este check-list deve ser preenchido pelo operador antes de iniciar as atividades e vistado pelo encarregado da área.")
    cell_obs.font = font_sub
    cell_obs.alignment = Alignment(horizontal="left", vertical="center")
    for c_idx in range(1, 35):
        ws.cell(row=row_atual, column=c_idx).border = borda_grade
    ws.row_dimensions[row_atual].height = 22
    row_atual += 1

    # 7. LEGENDA
    ws.merge_cells(start_row=row_atual, start_column=1, end_row=row_atual, end_column=34)
    cell_leg = ws.cell(row=row_atual, column=1, value="        LEGENDA:    O   - OK          X   - NG (Não Conforme)          N/A   - Não se Aplica")
    cell_leg.font = font_legenda
    cell_leg.alignment = Alignment(horizontal="left", vertical="center")
    for c_idx in range(1, 35):
        ws.cell(row=row_atual, column=c_idx).border = borda_grade
    ws.row_dimensions[row_atual].height = 25

    # Dimensionamento
    ws.column_dimensions['A'].width = 6    
    ws.column_dimensions['B'].width = 46   
    ws.column_dimensions['C'].width = 12   
    for i in range(4, 35):
        ws.column_dimensions[get_column_letter(i)].width = 3.0 

    ws.row_dimensions[1].height = 20
    ws.row_dimensions[2].height = 30 
    ws.row_dimensions[9].height = 28 

    # Salvar ou retornar
    if caminho_salvamento:
        wb.save(caminho_salvamento)
        return True
    else:
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

# ==============================================================================
# FUNÇÕES DO ATS (Word) - CORRIGIDAS PARA SUBSTITUIR TODAS AS TAGS
# ==============================================================================
def substituir_texto_completo(paragraph, mapeamento):
    """
    Substitui todas as tags em um parágrafo, incluindo texto dividido em múltiplos runs
    """
    # Junta todo o texto dos runs
    texto_completo = "".join([run.text for run in paragraph.runs])
    
    # Verifica se há tags para substituir
    houve_mudanca = False
    for tag, valor in mapeamento.items():
        if tag in texto_completo:
            texto_completo = texto_completo.replace(tag, str(valor))
            houve_mudanca = True
    
    # Se houve mudança, atualiza o parágrafo
    if houve_mudanca:
        # Limpa todos os runs existentes
        for run in paragraph.runs:
            run.text = ""
        # Adiciona o texto completo no primeiro run
        if paragraph.runs:
            paragraph.runs[0].text = texto_completo
        else:
            paragraph.add_run(texto_completo)
    
    return houve_mudanca

def substituir_tags_documento(doc, mapeamento):
    """
    Substitui todas as tags no documento inteiro (parágrafos e tabelas)
    """
    # Substituir em parágrafos normais
    for paragraph in doc.paragraphs:
        substituir_texto_completo(paragraph, mapeamento)
    
    # Substituir em todas as células de todas as tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    substituir_texto_completo(paragraph, mapeamento)

def processar_tabela_dinamica(doc, passos_atividade):
    """
    Processa a tabela dinâmica de tarefas/riscos/ações
    """
    for table in doc.tables:
        linha_modelo = None
        idx_tarefa = None
        idx_risco = None
        idx_acao = None
        
        # Procura pela linha com as tags especiais
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                texto_celula = "".join([run.text for run in cell.paragraphs[0].runs]) if cell.paragraphs else ""
                texto_celula = texto_celula.upper()
                
                if "{{TAREFAS}}" in texto_celula:
                    linha_modelo = row
                    idx_tarefa = idx
                if "{{RISCOS}}" in texto_celula:
                    linha_modelo = row
                    idx_risco = idx
                if "{{ACOES}}" in texto_celula:
                    linha_modelo = row
                    idx_acao = idx
            
            if linha_modelo is not None:
                break
        
        # Se encontrou a linha modelo, processa
        if linha_modelo is not None:
            passos_lista = passos_atividade.to_dict(orient="records")
            linha_referencia = linha_modelo
            
            for i, item in enumerate(passos_lista):
                tarefa = str(item.get("tarefa", "")).strip()
                risco = str(item.get("risco", "")).strip()
                acao = str(item.get("acoes", "")).strip()
                
                if tarefa.lower() == 'nan': tarefa = ""
                if risco.lower() == 'nan': risco = ""
                if acao.lower() == 'nan': acao = ""

                if i == 0:
                    nova_linha = linha_modelo
                else:
                    # Copia a linha modelo
                    tr_copiado = copy.deepcopy(linha_modelo._tr)
                    linha_referencia._tr.addnext(tr_copiado)
                    nova_linha = _Row(tr_copiado, table)
                    linha_referencia = nova_linha
                
                # Função para preencher célula
                def preencher_celula(celula, texto):
                    # Limpa a célula
                    for paragraph in celula.paragraphs:
                        for run in paragraph.runs:
                            run.text = ""
                    # Adiciona o novo texto
                    if celula.paragraphs:
                        p = celula.paragraphs[0]
                    else:
                        p = celula.add_paragraph()
                    run = p.add_run(texto)
                    run.font.size = Pt(7)
                
                if idx_tarefa is not None:
                    texto_tarefa = f"{i+1}. {tarefa}" if tarefa else ""
                    preencher_celula(nova_linha.cells[idx_tarefa], texto_tarefa)
                
                if idx_risco is not None:
                    texto_risco = f"• {risco}" if risco else ""
                    preencher_celula(nova_linha.cells[idx_risco], texto_risco)
                
                if idx_acao is not None:
                    texto_acao = f"• {acao}" if acao else ""
                    preencher_celula(nova_linha.cells[idx_acao], texto_acao)
            
            # Remove as tags da linha modelo após processar
            for cell in linha_modelo.cells:
                for paragraph in cell.paragraphs:
                    texto = "".join([run.text for run in paragraph.runs])
                    texto = texto.replace("{{TAREFAS}}", "").replace("{{RISCOS}}", "").replace("{{ACOES}}", "")
                    # Limpa e recoloca o texto
                    for run in paragraph.runs:
                        run.text = ""
                    if paragraph.runs:
                        paragraph.runs[0].text = texto
                    else:
                        paragraph.add_run(texto)

def gerar_documento_ats(dados, atividade_selecionada, passos):
    """
    Gera o documento ATS completo
    """
    template_path = "T.SHE.046 Safe Job Analyses.docx"
    
    if not os.path.exists(template_path):
        st.error(f"❌ Template não encontrado: {template_path}")
        return None

    try:
        # Carrega o documento
        doc = Document(template_path)
        
        # Mapeamento de substituições
        mapeamento = {
            "{{CONTRATADA}}": dados["contratada"].strip().upper(),
            "{{LOCAL}}": dados["local"].strip().upper(),
            "{{AREA}}": dados["area"].strip().upper(),
            "{{ÁREA}}": dados["area"].strip().upper(),
            "{{ATIVIDADE}}": atividade_selecionada.strip().upper(),
            "{{PROCESSO}}": dados["processo"].strip().upper()
        }
        
        # Substitui todas as tags no documento
        substituir_tags_documento(doc, mapeamento)
        
        # Processa a tabela dinâmica
        processar_tabela_dinamica(doc, passos)
        
        # Salva em memória
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"❌ Erro ao gerar documento: {str(e)}")
        return None

# ==============================================================================
# FUNÇÃO PARA CARREGAR DADOS DO EXCEL DO ATS
# ==============================================================================
@st.cache_data
def carregar_dados_ats():
    """Carrega os dados do arquivo Excel local banco_aprs.xlsx"""
    excel_path = "banco_aprs.xlsx"
    
    if os.path.exists(excel_path):
        try:
            df = pd.read_excel(excel_path)
            df.columns = df.columns.str.lower().str.strip()
            
            # Renomeia colunas para padronização
            rename_map = {
                'ações': 'acoes',
                'acao': 'acoes',
                'atividade': 'atividade',
                'tarefa': 'tarefa',
                'risco': 'risco',
                'área': 'area'
            }
            df.rename(columns=rename_map, inplace=True)
            
            # Verifica se as colunas necessárias existem
            colunas_necessarias = ['atividade', 'tarefa', 'risco', 'acoes']
            for col in colunas_necessarias:
                if col not in df.columns:
                    st.warning(f"⚠️ Coluna '{col}' não encontrada no Excel")
                    return pd.DataFrame(), []
            
            lista_atividades = df["atividade"].dropna().unique().tolist()
            return df, lista_atividades
            
        except Exception as e:
            st.error(f"❌ Erro ao ler o Excel: {str(e)}")
            return pd.DataFrame(), []
    
    return pd.DataFrame(), []

# ==============================================================================
# INTERFACE PRINCIPAL
# ==============================================================================
st.title("🏗️ Sistema Integrado de Segurança - SSMA")
st.markdown("---")

# Sidebar para navegação
st.sidebar.title("📋 Navegação")
modulo = st.sidebar.radio(
    "Selecione o Módulo:",
    ["📋 Checklist de Equipamentos", "📄 ATS - Análise de Trabalho Seguro"],
    index=0
)

# ==============================================================================
# MÓDULO 1: CHECKLIST DE EQUIPAMENTOS
# ==============================================================================
if modulo == "📋 Checklist de Equipamentos":
    st.header("📋 Gerador de Checklist de Equipamentos")
    st.markdown("Selecione um equipamento para gerar seu checklist de segurança.")
    
    # Verificar imagens
    st.sidebar.subheader("📸 Status das Imagens")
    imagens = {
        "Logo Heineken": r"C:\Users\dilceu.gomes\Desktop\check list\heineken.png",
        "Seta Verde": r"C:\Users\dilceu.gomes\Desktop\check list\seta.png",
        "SG": r"C:\Users\dilceu.gomes\Desktop\check list\SG.png"
    }
    
    for nome, caminho in imagens.items():
        if os.path.exists(caminho):
            st.sidebar.success(f"✅ {nome} encontrada")
        else:
            st.sidebar.warning(f"⚠️ {nome} não encontrada")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        equipamento_selecionado = st.selectbox(
            "Equipamento:",
            sorted(BANCO_EQUIPAMENTOS.keys())
        )
    
    with col2:
        st.write("")
        st.write("")
        if st.button("📊 Gerar Checklist", type="primary", use_container_width=True):
            with st.spinner("Gerando checklist..."):
                itens_checagem = BANCO_EQUIPAMENTOS[equipamento_selecionado]
                
                # Gerar Excel em memória
                excel_data = gerar_excel_checklist(
                    equipamento_selecionado,
                    itens_checagem,
                    None
                )
                
                if excel_data:
                    st.success("✅ Checklist gerado com sucesso!")
                    
                    # Botão de download
                    st.download_button(
                        label="📥 Baixar Checklist (Excel)",
                        data=excel_data,
                        file_name=f"Checklist_{equipamento_selecionado.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True
                    )
                    
                    # Preview dos itens
                    st.subheader("📝 Itens de Verificação:")
                    for idx, item in enumerate(itens_checagem, 1):
                        st.write(f"{idx}. {item}")

# ==============================================================================
# MÓDULO 2: ATS - ANÁLISE DE TRABALHO SEGURO
# ==============================================================================
else:
    st.header("📄 Gerador de ATS - Análise de Trabalho Seguro")
    
    # Verificar arquivos
    st.sidebar.subheader("📁 Status dos Arquivos")
    arquivos = {
        "Banco de Dados (banco_aprs.xlsx)": "banco_aprs.xlsx",
        "Template ATS (T.SHE.046 Safe Job Analyses.docx)": "T.SHE.046 Safe Job Analyses.docx"
    }
    
    for nome, caminho in arquivos.items():
        if os.path.exists(caminho):
            st.sidebar.success(f"✅ {nome} encontrado")
        else:
            st.sidebar.error(f"❌ {nome} não encontrado")
    
    # Carregar dados do Excel local
    df_excel, lista_atividades = carregar_dados_ats()
    
    if not df_excel.empty:
        st.success(f"✅ Banco de dados carregado com {len(lista_atividades)} atividades")
        
        col1, col2 = st.columns(2)
        with col1:
            atividade = st.selectbox("Atividade:", lista_atividades)
            processo = st.text_input("Processo:", "MANUTENÇÃO MECÂNICA")
        
        with col2:
            contratada = st.text_input("Contratada:", "MACROMAQ")
            local = st.text_input("Local:", "OFICINA")
            area = st.text_input("Área:", "EMPILHADEIRA")
        
        # Filtrar passos
        if atividade:
            passos = df_excel[df_excel["atividade"] == atividade]
            
            st.subheader("📋 Passos da Atividade")
            st.dataframe(passos, use_container_width=True)
            
            # Verificar template
            if os.path.exists("T.SHE.046 Safe Job Analyses.docx"):
                # Botão Gerar
                if st.button("🚀 Gerar ATS", type="primary", use_container_width=True):
                    if not passos.empty:
                        with st.spinner("Gerando documento ATS..."):
                            dados = {
                                "contratada": contratada,
                                "local": local,
                                "area": area,
                                "processo": processo
                            }
                            
                            resultado = gerar_documento_ats(
                                dados, 
                                atividade, 
                                passos
                            )
                            
                            if resultado:
                                st.success("✅ ATS gerado com sucesso!")
                                st.balloons()
                                st.download_button(
                                    "📥 Baixar ATS (Word)",
                                    resultado,
                                    file_name=f"ATS_{atividade.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True
                                )
                                
                                # Preview do conteúdo gerado
                                st.subheader("📄 Preview do Documento:")
                                st.info(f"""
                                **Contratada:** {contratada.upper()}
                                **Local:** {local.upper()}
                                **Área:** {area.upper()}
                                **Atividade:** {atividade.upper()}
                                **Processo:** {processo.upper()}
                                **Total de Passos:** {len(passos)}
                                """)
                            else:
                                st.error("❌ Erro ao gerar o documento. Verifique o template.")
                    else:
                        st.warning("⚠️ Nenhum passo encontrado para a atividade selecionada.")
            else:
                st.error("❌ Template do ATS não encontrado. Certifique-se que o arquivo 'T.SHE.046 Safe Job Analyses.docx' está na pasta.")
    else:
        st.error("❌ Banco de dados não encontrado. Certifique-se que o arquivo 'banco_aprs.xlsx' está na pasta.")
        st.info("""
        **Estrutura esperada do Excel:**
        - Coluna: `Atividade` (nome da atividade)
        - Coluna: `Tarefa` (tarefa específica)
        - Coluna: `Risco` (riscos associados)
        - Coluna: `Ações` (medidas de controle)
        
        **Exemplo:**
        | Atividade | Tarefa | Risco | Ações |
        |-----------|--------|-------|-------|
        | Manutenção | Desmontar | Queda | Usar cinto |
        """)

# ==============================================================================
# RODAPÉ
# ==============================================================================
st.markdown("---")
st.markdown(
    """
    <div style='text-align: center; color: #666; font-size: 14px;'>
    <b>Sistema Integrado de Segurança - SSMA</b> | Desenvolvido para gestão de segurança do trabalho
    </div>
    """,
    unsafe_allow_html=True
)